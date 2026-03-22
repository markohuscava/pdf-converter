import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import io
from PIL import Image

st.set_page_config(page_title="PDF Image Extractor", layout="wide")

def process_pdf(pdf_file, project_name):
    pdf_bytes = pdf_file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    word_doc = Document()
    
    # Nastavenie hlavičky
    section = word_doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Projekt: {project_name}"

    progress_bar = st.progress(0)
    status_text = st.empty()
    total_pages = len(doc)

    for page_num in range(total_pages):
        status_text.text(f"Spracovávam stranu {page_num + 1} z {total_pages}...")
        page = doc[page_num]
        image_list = page.get_images(full=True)
        
        if page_num > 0:
            word_doc.add_page_break()

        if image_list:
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                image_stream = io.BytesIO(image_bytes)
                
                # Vloženie obrázka (šírka nastavená na 6 palcov - cca šírka strany)
                try:
                    word_doc.add_picture(image_stream, width=Inches(6))
                except:
                    continue # Preskočiť poškodené obrázky
        
        progress_bar.progress((page_num + 1) / total_pages)

    status_text.text("Generujem DOCX súbor...")
    docx_output = io.BytesIO()
    word_doc.save(docx_output)
    docx_output.seek(0)
    return docx_output

# --- UI ---
st.title("🖼️ PDF Image to DOCX Converter")
st.info("Táto aplikácia odstráni text a ponechá len obrázky v novom Word dokumente.")

col1, col2 = st.columns([1, 1])

with col1:
    project_name = st.text_input("Názov projektu (bude v hlavičke):")
    uploaded_file = st.file_uploader("Nahrajte PDF súbor", type="pdf")

if uploaded_file and project_name:
    # Preview
    with col2:
        st.subheader("Náhľad (prvé 3 strany)")
        doc_preview = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
        preview_limit = min(3, len(doc_preview))
        for i in range(preview_limit):
            page = doc_preview[i]
            pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5)) # menšie rozlíšenie pre rýchly náhľad
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            st.image(img, caption=f"Strana {i+1}")

    if st.button("🚀 Spustiť konverziu"):
        result_docx = process_pdf(uploaded_file, project_name)
        
        st.success("Hotovo!")
        st.download_button(
            label="⬇️ Stiahnuť DOCX",
            data=result_docx,
            file_name=f"{project_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )